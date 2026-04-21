import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_FILES = [
    BASE_DIR / "fib_200.json",
    BASE_DIR / "fib_200_generated.json",
]


def load_questions(json_path: Path):
    with json_path.open("r", encoding="utf-8") as file_handle:
        payload = json.load(file_handle)
    return payload["meta"], payload["questions"]


def unique_answers(questions):
    answers = []
    seen = set()
    for question in questions:
        answer = str(question["answer"]).strip()
        key = answer.casefold()
        if key in seen:
            continue
        seen.add(key)
        answers.append(answer)
    return answers


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def build_document(meta, questions, output_path: Path):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(meta["course"])
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Fill-in-the-Blank Question Set")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    meta_lines = [
        f'Type: {meta["type"]}',
        f'Total questions: {meta["total_questions"]}',
        f'Lectures covered: {", ".join(str(item) for item in meta["lectures_covered"])}',
        f'Blank token: {meta["blank_token"]}',
    ]
    for line in meta_lines:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(line)

    document.add_paragraph()

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["#", "Lecture", "Topic", "Question", "Answer"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)

    for question in questions:
        row = table.add_row().cells
        set_cell_text(row[0], str(question["id"]))
        set_cell_text(row[1], str(question["lecture"]))
        set_cell_text(row[2], str(question["topic"]))
        set_cell_text(row[3], str(question["question"]))
        set_cell_text(row[4], str(question["answer"]))

    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("Vocabulary Bank")
    heading_run.bold = True
    heading_run.font.size = Pt(14)

    vocab_intro = document.add_paragraph()
    vocab_intro.paragraph_format.space_after = Pt(6)
    vocab_intro.add_run("Unique answer entries listed in order of appearance.")

    for answer in unique_answers(questions):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(answer)

    document.save(output_path)


def write_vocab_file(questions, output_path: Path):
    answers = unique_answers(questions)
    with output_path.open("w", encoding="utf-8", newline="\n") as file_handle:
        file_handle.write("Vocabulary Bank\n")
        file_handle.write("\n")
        for index, answer in enumerate(answers, 1):
            file_handle.write(f"{index}. {answer}\n")


def main():
    for source_path in SOURCE_FILES:
        if not source_path.exists():
            raise FileNotFoundError(source_path)

        meta, questions = load_questions(source_path)
        output_stem = source_path.with_suffix("")
        docx_path = output_stem.with_suffix(".docx")
        vocab_path = output_stem.with_name(f"{output_stem.name}_vocabulary_bank.txt")

        build_document(meta, questions, docx_path)
        write_vocab_file(questions, vocab_path)
        print(f"Wrote {docx_path.name} and {vocab_path.name}")


if __name__ == "__main__":
    main()